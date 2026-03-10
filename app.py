"""
DealMemo AI — v4.0
Professional IC Memo Generator — Indian Startup Ecosystem
Fixes: markdown rendering in Word, company pre-analysis, industry-dynamic prompts
"""

import os, json, time, re
import streamlit as st
from groq import Groq
import PyPDF2
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

# ─── SECRETS — works on Streamlit Cloud AND local .env ───────────────────────
def get_secret(key):
    try:
        return st.secrets[key]
    except Exception:
        pass
    return os.environ.get(key, "")

# ─── COLOURS ─────────────────────────────────────────────────────────────────
NAVY       = "1F3864"
DARK_BLUE  = "1F4E79"
MID_BLUE   = "2E75B6"
LIGHT_BLUE = "D6E4F0"
GREEN      = "1E8449"
AMBER      = "D68910"
RED        = "C0392B"
GREY       = "7F8C8D"
WHITE      = "FFFFFF"

# ─── GROQ CLIENT ─────────────────────────────────────────────────────────────
@st.cache_resource
def get_client():
    key = get_secret("GROQ_API_KEY")
    if not key:
        st.error(
            "GROQ_API_KEY not found. "
            "Go to Streamlit Cloud -> Settings -> Secrets and add:\n"
            'GROQ_API_KEY = "gsk_your_key_here"'
        )
        st.stop()
    return Groq(api_key=key)

PRIMARY_MODEL  = "llama-3.3-70b-versatile"
FALLBACK_MODEL = "llama3-70b-8192"

# ─── SAFE AI CALL ─────────────────────────────────────────────────────────────
def safe_ai_call(messages, expect_json=False, max_tokens=1200):
    client = get_client()
    for attempt, model in enumerate([PRIMARY_MODEL, FALLBACK_MODEL]):
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=0.15 if expect_json else 0.3,
                max_tokens=max_tokens,
            )
            raw = resp.choices[0].message.content.strip()
            if expect_json:
                raw = re.sub(r"```(?:json)?", "", raw).strip().strip("`").strip()
                try:
                    json.loads(raw)
                    return raw
                except json.JSONDecodeError:
                    m = re.search(r"\{[\s\S]*\}", raw)
                    if m:
                        try:
                            json.loads(m.group(0))
                            return m.group(0)
                        except Exception:
                            pass
                return "{}"
            return raw
        except Exception as e:
            err = str(e).lower()
            if "rate" in err or "429" in err or "quota" in err:
                time.sleep(12 if attempt == 0 else 20)
                continue
            return (
                '{"error": "API error"}' if expect_json
                else f"[Generation error: {str(e)[:60]}. Please retry.]"
            )
    return (
        '{"error": "Rate limit"}' if expect_json
        else "[Rate limit reached. Please wait 60 seconds and retry.]"
    )

# ─── PDF EXTRACTION ───────────────────────────────────────────────────────────
def extract_pdf(pdf_file):
    try:
        reader = PyPDF2.PdfReader(pdf_file)
        pages  = [p.extract_text() for p in reader.pages
                  if p.extract_text() and p.extract_text().strip()]
        return "\n\n".join(pages)[:14000]
    except Exception:
        return ""

# ════════════════════════════════════════════════════════════════════════════
#  PRE-ANALYSIS — Detects industry, stage, listing status
#  This context is injected into EVERY subsequent prompt
# ════════════════════════════════════════════════════════════════════════════

def analyse_company_profile(deck, company):
    raw = safe_ai_call(
        messages=[
            {"role": "system",
             "content": "You are a data extractor. Return only valid JSON. No text outside JSON."},
            {"role": "user", "content": f"""
Analyse this pitch deck for {company} and return a company profile.
Return ONLY this JSON:

{{
  "company_stage": "Pre-revenue / Pre-seed / Seed / Series A / Series B / Series C+ / Growth / Public / PE Buyout",
  "listing_status": "Private / Public / Recently IPO'd",
  "stock_exchange": "NSE / BSE / NASDAQ / NYSE / Not Listed",
  "primary_industry": "Fintech / Edtech / Healthtech / D2C / SaaS / Marketplace / Logistics / Agritech / Proptech / Consumer Internet / Gaming / Deep Tech / Media / Other",
  "sub_sector": "Specific sub-sector in 3-5 words",
  "business_model": "B2B / B2C / B2B2C / Marketplace / SaaS / D2C / Hybrid",
  "revenue_status": "No revenue / Pre-revenue / Early revenue / Scaling revenue / Profitable",
  "india_focus": "India-only / India-primary / India + Southeast Asia / Global",
  "relevant_indian_benchmarks": ["Company1", "Company2", "Company3"],
  "gross_margin_benchmark": "e.g. 15-25% for food delivery / 60-80% for SaaS",
  "cac_ltv_relevant": "Yes / No",
  "regulatory_bodies": ["SEBI", "RBI"],
  "known_public_data_available": "Yes / No",
  "stage_summary": "One sentence describing what stage analysis lens to apply"
}}

Deck content:
{deck[:3000]}
"""}
        ],
        expect_json=True,
        max_tokens=600
    )
    try:
        profile = json.loads(raw)
        if profile.get("company_stage"):
            return profile
    except Exception:
        pass
    return {
        "company_stage": "Early Stage",
        "listing_status": "Private",
        "stock_exchange": "Not Listed",
        "primary_industry": "Consumer Internet",
        "sub_sector": "General",
        "business_model": "B2C",
        "revenue_status": "Early revenue",
        "india_focus": "India-primary",
        "relevant_indian_benchmarks": [],
        "gross_margin_benchmark": "Sector-specific",
        "cac_ltv_relevant": "Yes",
        "regulatory_bodies": ["SEBI"],
        "known_public_data_available": "No",
        "stage_summary": "Early-stage company. Use conservative benchmarks."
    }

def profile_to_context(profile, company):
    is_public = profile.get("known_public_data_available", "No")
    return f"""
COMPANY PROFILE (pre-analysed - calibrate everything to this):
- Company: {company}
- Listing Status: {profile.get('listing_status','Unknown')} ({profile.get('stock_exchange','N/A')})
- Stage: {profile.get('company_stage','Unknown')}
- Industry: {profile.get('primary_industry','Unknown')} / {profile.get('sub_sector','')}
- Business Model: {profile.get('business_model','Unknown')}
- Revenue Status: {profile.get('revenue_status','Unknown')}
- Gross Margin Benchmark for this sector: {profile.get('gross_margin_benchmark','Sector-specific')}
- CAC/LTV Relevant: {profile.get('cac_ltv_relevant','Yes')}
- Key Regulators: {', '.join(profile.get('regulatory_bodies', ['SEBI']))}
- Public Data Available: {is_public}
- Indian Comparables: {', '.join(profile.get('relevant_indian_benchmarks', []))}
- Analysis Lens: {profile.get('stage_summary','')}

{"CRITICAL: This is a PUBLIC company. You MUST use your training knowledge of " + company + " actual disclosed financials. Do NOT say NOT PROVIDED for publicly available metrics." if is_public == "Yes" else "This is a private company. Extract from deck. Use NOT PROVIDED only if genuinely unavailable."}

Calibrate ALL benchmarks to the {profile.get('primary_industry','relevant')} sector in India at {profile.get('company_stage','')} stage.
"""

# ════════════════════════════════════════════════════════════════════════════
#  MASTER SYSTEM PROMPT
# ════════════════════════════════════════════════════════════════════════════

SYSTEM = """You are Arjun Mehta, Principal at Sequoia Capital India with 9 years
experience. You have personally written IC memos for 340+ deals. Three unicorns
in your portfolio.

Your IC memos are known for:
1. Brutal honesty - you flag every red flag, even when the founder is in the room
2. India-first lens - INR figures, Indian benchmarks, Indian regulatory context
3. Decisive conclusions - you make a call, you never hedge
4. Real data - for public companies you cite actual disclosed metrics

Writing rules (CRITICAL - follow exactly):
- Short punchy sentences. No corporate filler words.
- Numbers over adjectives. Always.
- Reference real Indian companies as benchmarks by name.
- Write subheading labels as plain text on their own line, NO asterisks.
- Do NOT use ** or __ or any markdown symbols anywhere in your response.
- The output goes directly into a Word document. Clean text only.
- If something is truly unknown: [NOT PROVIDED - Request in management Q&A]
- Never fabricate financial figures."""

# ════════════════════════════════════════════════════════════════════════════
#  SECTION PROMPTS
# ════════════════════════════════════════════════════════════════════════════

def prompt_exec(deck, company, fund_type, ctx):
    return f"""
{ctx}

Write the EXECUTIVE SUMMARY AND INVESTMENT THESIS for {company}.
Memo format: {fund_type}

Pitch deck content:
---
{deck[:5000]}
---

Write with these subheading labels on separate lines (NO asterisks, clean text):

The Opportunity
[2-3 sentences. Specific problem, who suffers, scale of pain. Lead with a number.]

What {company} Does
[Plain English product description. One paragraph. How it works end-to-end.]

Why This, Why Now
[Specific India tailwinds. Regulatory shift, tech unlock, consumer behaviour change.
Be specific - not "India is a large market".]

Traction Snapshot
[4-5 most important metrics. For public companies use actual disclosed figures.
Format: Metric Name: Value]

Preliminary View
[One sentence: PROCEED TO DILIGENCE / CONDITIONAL / PASS
Then: Conviction level: High/Medium/Low - because [specific reason].]

Maximum 300 words. Write like you have 5 minutes before IC starts.
NO asterisks. NO markdown. Clean text only.
"""

def prompt_overview(deck, company, fund_type, ctx):
    return f"""
{ctx}

Write the COMPANY OVERVIEW AND BUSINESS MODEL for {company}.
Memo format: {fund_type}

Pitch deck:
---
{deck[:5000]}
---

Subheadings as plain text labels (NO asterisks):

What They Do
[2-3 sentences. Describe it to a sceptical LP, not using the company's own marketing language.]

How They Make Money
[All revenue streams. Pricing. Take rate if marketplace. Proven or theoretical?]

Customer Profile
[Exact customer segment. What they pay and why. Geography.]

Current Stage and Validation
[What has been proven vs what is still assumption.]

Regulatory and Compliance
[Relevant Indian regulatory bodies specific to this industry. Flag compliance risks explicitly.]

250-280 words. Short paragraphs. No bullets. No markdown.
"""

def prompt_market(deck, company, fund_type, ctx):
    return f"""
{ctx}

Write the MARKET OPPORTUNITY section for {company}.
Memo format: {fund_type}

Pitch deck:
---
{deck[:5000]}
---

Subheadings as plain text labels (NO asterisks):

Market Context
[Why this market is interesting in India specifically. One structural shift.
Cite a specific data point - not generic India statements.]

Size Assessment
[TAM, SAM, SOM in INR Crores. Your own bottom-up estimate.
Challenge inflated claims: "The company claims X. Our estimate is Y because Z."]

Market Structure
[Fragmented or consolidated. Top 2-3 players. How much is still offline.]

Analyst View
[Can this company realistically build Rs.500 Cr+ revenue? What market share?
Your honest call - not a hedge.]

250-280 words. No bullets. No markdown.
"""

def prompt_recommendation(deck, company, fund_type, ctx):
    return f"""
{ctx}

Write the INVESTMENT RECOMMENDATION for {company}.
This is the most important section. Be decisive. Memo format: {fund_type}

Pitch deck:
---
{deck[:5000]}
---

Subheadings as plain text labels (NO asterisks):

Recommendation
[First line must be exactly one of:
RECOMMENDATION: PROCEED TO FULL DILIGENCE
RECOMMENDATION: CONDITIONAL PROCEED
RECOMMENDATION: PASS
Then state check size and ownership target, OR conditions, OR reason for pass.]

Bull Case - Three Year View
[If everything goes right. Revenue, market position, comparable exit.
Expected MOIC. Reference a DIFFERENT company that succeeded in this space as comparable.
Do NOT reference {company} itself as the benchmark.]

Bear Case
[Most realistic failure path. Not a black swan - the most likely path to zero.
Probability estimate as a percentage.]

Critical Diligence Items
[Exactly 5 questions specific to THIS deal. Not generic. Questions that address
the specific risks you identified above.]

Proposed Terms
[Information rights, board seat, pro-rata, any milestone tranches.
Stage-appropriate for {company}.]

320-360 words. The IC reads this first. Make the recommendation impossible to miss.
NO asterisks. NO markdown. Clean text only.
"""

def prompt_competitive(deck, company, profile):
    industry = profile.get("primary_industry", "the relevant sector")
    stage    = profile.get("company_stage", "")
    return f"""
Analyse the competitive landscape for {company}.
Industry: {industry}. Stage: {stage}.

Pitch deck:
---
{deck[:4000]}
---

Return ONLY valid JSON. No text outside JSON. No markdown fences.

{{
  "rows": [
    {{
      "company": "Competitor name",
      "stage": "Series B / Public / Bootstrapped etc",
      "funding_raised": "Total raised e.g. $45M or Rs.320 Cr or Undisclosed",
      "core_product": "6 words max",
      "revenue_model": "5 words max",
      "india_focus": "Primary / Secondary / None",
      "moat": "Single strongest competitive advantage",
      "weakness": "Most exploitable weakness",
      "threat_to_subject": "High / Medium / Low",
      "threat_rationale": "One sentence why"
    }}
  ],
  "competitive_summary": "2-sentence analyst view on competitive intensity and {company} positioning"
}}

Rules:
- 4-5 REAL competitors from the {industry} space in India using your training knowledge
- Last row must be {company} with company field = "{company} (Subject)" and threat_to_subject = "Subject"
- Prioritise India-based competitors
- Be analytically honest - do not favour {company}
- All strings. Undisclosed for unknown funding. No nulls.
"""

def prompt_team(deck, company, profile):
    industry = profile.get("primary_industry", "")
    stage    = profile.get("company_stage", "")
    return f"""
Assess the founding and leadership team of {company}.
Industry: {industry}. Stage: {stage}.

Pitch deck:
---
{deck[:4000]}
---

Return ONLY valid JSON. No text outside JSON. No markdown fences.

{{
  "members": [
    {{
      "name": "Full name or Not Disclosed",
      "title": "Job title",
      "previous_employer": "Most impressive prior company",
      "previous_role": "Prior role title",
      "years_experience": "e.g. 8 years",
      "direct_relevance": "How prior experience applies to {company} specifically",
      "education": "Degree and institution",
      "flag": "Green / Yellow / Red",
      "flag_note": "One specific sentence explaining this flag"
    }}
  ],
  "team_verdict": "3-sentence honest assessment. Name specific gaps.",
  "missing_hires": "Critical C-suite or functional roles not yet filled",
  "founder_risk": "High / Medium / Low",
  "founder_risk_reason": "One sentence"
}}

Green = proven operator, directly relevant to {industry}, strong pedigree
Yellow = adequate but a specific named gap exists - name it
Red = no relevant industry experience, undisclosed key person, or first-time founder in regulated space
Add row with name = "[Role] Not Disclosed" and flag = Red for missing C-suite roles.
All strings. No nulls.
"""

def prompt_financials(deck, company, fund_type, profile):
    industry  = profile.get("primary_industry", "")
    stage     = profile.get("company_stage", "")
    gm_bench  = profile.get("gross_margin_benchmark", "Sector-specific")
    is_public = profile.get("known_public_data_available", "No")
    cac_ltv   = profile.get("cac_ltv_relevant", "Yes")
    biz_model = profile.get("business_model", "")
    return f"""
Extract and assess unit economics for {company}.
Industry: {industry}. Stage: {stage}. Model: {biz_model}.
Memo format: {fund_type}.
Gross margin benchmark for {industry}: {gm_bench}
Public financial data available: {is_public}

{"CRITICAL INSTRUCTION: This is a PUBLIC company. Use your training knowledge of " + company + " actual reported financials from their public disclosures and earnings reports. Do NOT say NOT PROVIDED for publicly known metrics." if is_public == "Yes" else "Private company - extract from deck. NOT PROVIDED only if truly absent."}

Pitch deck:
---
{deck[:4000]}
---

Return ONLY valid JSON. No text outside JSON. No markdown fences.

{{
  "metrics": [
    {{
      "metric": "Metric name",
      "reported_value": "Value from deck or public disclosure or NOT PROVIDED",
      "analyst_comment": "One sentence assessment specific to {industry} at {stage} stage",
      "india_benchmark": "Typical range for {industry} companies in India at {stage} stage",
      "rating": "Strong / Acceptable / Weak / Not Provided"
    }}
  ],
  "overall_financial_health": "Strong / Adequate / Concerning / Insufficient Data",
  "headline_concern": "Single most important financial red flag - be specific to {company}",
  "headline_positive": "Single most compelling financial positive - be specific",
  "burn_runway_comment": "Burn efficiency and runway assessment calibrated to {stage} stage"
}}

Include these metrics calibrated to {industry} (use actual figures for public companies):
1. Revenue Run Rate (use ARR/MRR/GMV - whichever fits {industry})
2. Revenue Growth Rate (MoM or YoY as appropriate for stage)
3. Gross Margin % - benchmark: {gm_bench}
4. Contribution Margin
5. {"CAC - Customer Acquisition Cost" if "Yes" in cac_ltv else "Key Volume Metric for " + industry}
6. {"LTV - Customer Lifetime Value" if "Yes" in cac_ltv else "Revenue per Customer"}
7. {"LTV / CAC Ratio" if "Yes" in cac_ltv else "Efficiency Metric"}
8. Payback Period
9. Monthly Burn Rate
10. Runway in months
11. {"NRR or Churn Rate" if "SaaS" in industry or "B2B" in biz_model else "Monthly Active Users or Customers"}
12. {"MRR or ARR" if "SaaS" in industry else "Order Volume or GMV"}

Rating: Strong = above {industry} benchmarks / Acceptable = within range / Weak = below benchmarks
"""

def prompt_risks(deck, company, fund_type, profile):
    industry   = profile.get("primary_industry", "")
    stage      = profile.get("company_stage", "")
    reg_bodies = ", ".join(profile.get("regulatory_bodies", ["SEBI"]))
    return f"""
Identify investment risks for {company}.
Industry: {industry}. Stage: {stage}.
Key regulators: {reg_bodies}.

Pitch deck:
---
{deck[:4000]}
---

Return ONLY valid JSON. No text outside JSON. No markdown fences.

{{
  "risks": [
    {{
      "id": "R1",
      "category": "Market / Execution / Team / Regulatory / Financial / Technology / Competition",
      "title": "4-6 words specific to {company}",
      "description": "2 sentences specific to {company} and {industry}. Not generic.",
      "probability": "High / Medium / Low",
      "impact": "High / Medium / Low",
      "time_horizon": "0-12 months / 1-3 years / 3+ years",
      "mitigant": "Specific mitigating factor or action",
      "residual_risk": "High / Medium / Low",
      "deal_breaker": "Yes / No"
    }}
  ],
  "risk_summary": "2 sentences. Overall risk profile and the single risk that could kill this deal.",
  "overall_risk_rating": "High / Medium-High / Medium / Medium-Low / Low"
}}

Exactly 6 risks:
R1 = Market risk specific to {industry} in India
R2 = Execution risk specific to {company} operating model
R3 = Team or founder risk - be honest even if team looks strong
R4 = Regulatory risk involving {reg_bodies} specific to {industry}
R5 = Competitive risk - name ONE specific competitor that threatens {company}
R6 = Financial risk for {stage} stage company

deal_breaker = Yes only if this single risk alone would cause you to pass.
All strings. No nulls.
"""

def prompt_tam(deck, company, profile):
    industry = profile.get("primary_industry", "")
    india    = profile.get("india_focus", "India")
    return f"""
Extract or estimate TAM/SAM/SOM for {company} in INR Crores.
Industry: {industry}. Market: {india}.

Return ONLY this JSON. No text. No markdown:
{{"tam": <integer INR Crores>, "sam": <integer INR Crores>, "som": <integer INR Crores>}}

SAM must be less than TAM. SOM must be less than SAM. Integers only.
Use your knowledge of the {industry} market in India.

Deck: {deck[:2000]}
"""

# ════════════════════════════════════════════════════════════════════════════
#  WORD DOCUMENT HELPERS
# ════════════════════════════════════════════════════════════════════════════

def hex_rgb(h):
    h = h.lstrip("#")
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def cell_bg(cell, color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"), color.lstrip("#")); tcPr.append(shd)

def cell_pad(cell, t=80, b=80, l=120, r=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top",t),("bottom",b),("left",l),("right",r)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val)); el.set(qn("w:type"),"dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

def section_head(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text.upper())
    r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = hex_rgb(NAVY); r.font.name = "Calibri"
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"8")
    bot.set(qn("w:space"),"2"); bot.set(qn("w:color"), MID_BLUE)
    pBdr.append(bot); pPr.append(pBdr)

def add_rich_text(doc, text):
    """
    Converts AI text to properly formatted Word paragraphs.
    Strips all ** markers. Detects subheadings. Never shows asterisks.
    """
    if not text:
        return
    lines = text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            continue

        clean = re.sub(r"\*+", "", stripped).strip()
        if not clean:
            continue

        original_had_bold = re.search(r"\*\*", stripped) is not None
        is_short_cap = (
            len(clean) < 70
            and re.match(r"^[A-Z]", clean)
            and not clean.endswith(",")
            and clean.count(".") <= 1
        )
        is_subhead = original_had_bold or (
            is_short_cap and not any(
                clean.lower().startswith(kw)
                for kw in ["the ", "a ", "an ", "this ", "in ", "for ",
                           "by ", "with ", "at ", "if ", "we ", "our ",
                           "note", "source"]
            )
        )

        if is_subhead:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(9)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(clean)
            r.bold = True; r.font.size = Pt(10)
            r.font.color.rgb = hex_rgb(DARK_BLUE)
            r.font.name = "Calibri"
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(5)
            p.paragraph_format.space_before = Pt(1)
            r = p.add_run(clean)
            r.font.size = Pt(10)
            r.font.name = "Calibri"

def styled_table(doc, headers, rows, alt=LIGHT_BLUE):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    auto_color = {
        "High": RED, "Medium": AMBER, "Low": GREEN,
        "Strong": GREEN, "Acceptable": AMBER, "Weak": RED,
        "Not Provided": GREY, "Not Disclosed": GREY,
        "Green": GREEN, "Yellow": AMBER, "Red": RED,
        "Yes": RED, "No": GREEN,
        "Primary": GREEN, "Secondary": AMBER, "None": GREY,
        "Subject": MID_BLUE, "Medium-High": AMBER, "Medium-Low": GREEN,
        "Profitable": GREEN, "Concerning": RED, "Adequate": AMBER,
    }

    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        cell_bg(c, DARK_BLUE); cell_pad(c)
        c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(8.5)
        r.font.color.rgb = hex_rgb(WHITE); r.font.name = "Calibri"
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for ri, row_vals in enumerate(rows):
        row = t.add_row()
        bg  = alt if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row_vals):
            c = row.cells[ci]
            cell_pad(c)
            sv = str(val)
            override = auto_color.get(sv)
            if override:
                cell_bg(c, override)
                c.paragraphs[0].clear()
                r = c.paragraphs[0].add_run(sv)
                r.bold = True; r.font.size = Pt(8.5)
                r.font.color.rgb = hex_rgb(WHITE); r.font.name = "Calibri"
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                cell_bg(c, bg)
                c.paragraphs[0].clear()
                r = c.paragraphs[0].add_run(sv)
                r.font.size = Pt(8.5); r.font.name = "Calibri"
    doc.add_paragraph()

def embed_img(doc, buf, width=6.0, caption=None):
    buf.seek(0)
    doc.add_picture(buf, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cp.runs:
            r.font.size = Pt(8); r.font.color.rgb = hex_rgb(GREY)
            r.font.name = "Calibri"
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
#  CHARTS
# ════════════════════════════════════════════════════════════════════════════

def chart_tam(tam, company):
    fig, ax = plt.subplots(figsize=(7, 3.2))
    fig.patch.set_facecolor("#F7FAFC"); ax.set_facecolor("#F7FAFC")
    labels = ["TAM", "SAM", "SOM"]
    vals   = [tam.get("tam",45000), tam.get("sam",10000), tam.get("som",2000)]
    cols   = ["#1F4E79", "#2E75B6", "#9DC3E6"]
    bars   = ax.barh(labels, vals, color=cols, height=0.45,
                     edgecolor="white", linewidth=1.5)
    for bar, v in zip(bars, vals):
        ax.text(bar.get_width() + max(vals)*0.015,
                bar.get_y() + bar.get_height()/2,
                f"Rs.{v:,} Cr", va="center", fontsize=10,
                color="#1F3864", fontweight="bold")
    ax.set_title(f"{company} - Market Sizing (INR Crores)",
                 fontsize=11, fontweight="bold", color="#1F3864", pad=12)
    ax.set_xlabel("INR Crores", fontsize=9, color="#7F8C8D")
    ax.set_xlim(0, max(vals)*1.28)
    ax.tick_params(labelsize=10, colors="#34495E")
    for s in ["top","right"]: ax.spines[s].set_visible(False)
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format="png", dpi=150,
                bbox_inches="tight", facecolor="#F7FAFC")
    plt.close(fig); buf.seek(0)
    return buf

def chart_risk(risks):
    fig, ax = plt.subplots(figsize=(6.5, 4.5))
    fig.patch.set_facecolor("#F7FAFC"); ax.set_facecolor("#FDFEFE")
    zones = [
        (0,0,1,1,"#EBF5FB"),(1,0,1,1,"#FEF9E7"),(2,0,1,1,"#FDEDEC"),
        (0,1,1,1,"#FEF9E7"),(1,1,1,1,"#FEF5E7"),(2,1,1,1,"#FDEDEC"),
        (0,2,1,1,"#FDEDEC"),(1,2,1,1,"#FDEDEC"),(2,2,1,1,"#F9EBEA"),
    ]
    for x,y,w,h,c in zones:
        ax.add_patch(mpatches.Rectangle((x,y),w,h,color=c,zorder=0))
    lv = {"Low":0, "Medium":1, "High":2}
    cl = {"High":"#C0392B", "Medium":"#D68910", "Low":"#1E8449"}
    plotted = []
    for r in risks[:6]:
        px = lv.get(r.get("probability","Medium"),1) + 0.5
        py = lv.get(r.get("impact","Medium"),1) + 0.5
        jit = 0.1*sum(1 for p in plotted if abs(p[0]-px)<0.25 and abs(p[1]-py)<0.25)
        px += jit; py += jit; plotted.append((px,py))
        col = cl.get(r.get("probability","Medium"),"#7F8C8D")
        db  = "* " if r.get("deal_breaker")=="Yes" else ""
        ax.scatter(px,py,s=220,color=col,zorder=5,edgecolors="white",linewidths=2)
        ax.annotate(
            db+r.get("id","")+" "+r.get("title","")[:20],
            (px,py), textcoords="offset points", xytext=(8,5),
            fontsize=7.5, color="#1F3864", fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.25",facecolor="white",
                      edgecolor="#BDC3C7",alpha=0.9))
    ax.set_xlim(0,3); ax.set_ylim(0,3)
    ax.set_xticks([0.5,1.5,2.5]); ax.set_yticks([0.5,1.5,2.5])
    ax.set_xticklabels(["Low","Medium","High"],fontsize=10,color="#5D6D7E")
    ax.set_yticklabels(["Low","Medium","High"],fontsize=10,color="#5D6D7E")
    ax.set_xlabel("Probability",fontsize=9,color="#7F8C8D",labelpad=8)
    ax.set_ylabel("Impact",fontsize=9,color="#7F8C8D",labelpad=8)
    ax.set_title("Risk Matrix  (* = Deal Breaker Risk)",
                 fontsize=10,fontweight="bold",color="#1F3864",pad=12)
    ax.grid(True,linestyle=":",alpha=0.4,color="#BDC3C7")
    for s in ["top","right"]: ax.spines[s].set_visible(False)
    patches = [
        mpatches.Patch(color="#C0392B",label="High probability"),
        mpatches.Patch(color="#D68910",label="Medium probability"),
        mpatches.Patch(color="#1E8449",label="Low probability"),
    ]
    ax.legend(handles=patches,fontsize=7.5,loc="upper left",framealpha=0.85)
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf,format="png",dpi=150,bbox_inches="tight",facecolor="#F7FAFC")
    plt.close(fig); buf.seek(0)
    return buf

def chart_financials(metrics):
    rc = {"Strong":"#1E8449","Acceptable":"#D68910",
          "Weak":"#C0392B","Not Provided":"#95A5A6"}
    names  = [m.get("metric","")[:32] for m in metrics[:12]]
    colors = [rc.get(m.get("rating","Not Provided"),"#95A5A6") for m in metrics[:12]]
    fig, ax = plt.subplots(figsize=(7.5, len(names)*0.5+1.2))
    fig.patch.set_facecolor("#F7FAFC"); ax.set_facecolor("#F7FAFC")
    bars = ax.barh(names,[1]*len(names),color=colors,height=0.62,
                   edgecolor="white",linewidth=0.8)
    for bar, m in zip(bars,metrics[:12]):
        val    = m.get("reported_value","N/A")[:28]
        rating = m.get("rating","N/A")
        ax.text(0.02, bar.get_y()+bar.get_height()/2,
                f"  {val}   [{rating}]",
                va="center",fontsize=8.5,color="white",fontweight="bold")
    ax.set_xlim(0,1); ax.set_xticks([])
    ax.set_title("Unit Economics Dashboard",
                 fontsize=11,fontweight="bold",color="#1F3864",pad=12)
    ax.tick_params(colors="#34495E",labelsize=8.5)
    for s in ["top","right","bottom"]: ax.spines[s].set_visible(False)
    patches = [
        mpatches.Patch(color="#1E8449",label="Strong"),
        mpatches.Patch(color="#D68910",label="Acceptable"),
        mpatches.Patch(color="#C0392B",label="Weak"),
        mpatches.Patch(color="#95A5A6",label="Not Provided"),
    ]
    ax.legend(handles=patches,fontsize=8,loc="lower right",framealpha=0.85)
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf,format="png",dpi=150,bbox_inches="tight",facecolor="#F7FAFC")
    plt.close(fig); buf.seek(0)
    return buf

# ════════════════════════════════════════════════════════════════════════════
#  WORD DOCUMENT BUILDER
# ════════════════════════════════════════════════════════════════════════════

def build_memo(texts, tables, charts, company, fund_type, profile):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin   = Inches(1.0)
    sec.right_margin  = Inches(1.0)

    for text, size, color, bold in [
        ("STRICTLY CONFIDENTIAL  -  FOR IC USE ONLY", 8, GREY, False),
        ("INVESTMENT COMMITTEE MEMORANDUM", 20, NAVY, True),
        (company.upper(), 16, MID_BLUE, True),
        (f"{profile.get('primary_industry','')}  |  "
         f"{profile.get('company_stage','')}  |  "
         f"{profile.get('listing_status','Private')}", 9, GREY, False),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.bold = bold
        r.font.size = Pt(size); r.font.name = "Calibri"
        r.font.color.rgb = hex_rgb(color)

    doc.add_paragraph()
    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    for i, (k, v) in enumerate([
        ("Fund Format",   fund_type),
        ("Industry",      f"{profile.get('primary_industry','')} - {profile.get('sub_sector','')}"),
        ("Company Stage", profile.get("company_stage", "")),
        ("Prepared By",   "DealMemo AI  |  First-Draft Analyst Review"),
    ]):
        meta.cell(i,0).text = k; meta.cell(i,1).text = v
        cell_bg(meta.cell(i,0), LIGHT_BLUE)
        cell_pad(meta.cell(i,0)); cell_pad(meta.cell(i,1))
        for c in [meta.cell(i,0), meta.cell(i,1)]:
            for r2 in c.paragraphs[0].runs:
                r2.font.size = Pt(9); r2.font.name = "Calibri"
        meta.cell(i,0).paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    section_head(doc, "1.  Executive Summary & Investment Thesis")
    add_rich_text(doc, texts.get("exec",""))
    doc.add_paragraph()

    section_head(doc, "2.  Company Overview & Business Model")
    add_rich_text(doc, texts.get("overview",""))
    doc.add_paragraph()

    section_head(doc, "3.  Market Opportunity")
    add_rich_text(doc, texts.get("market",""))
    if "tam" in charts:
        doc.add_paragraph()
        embed_img(doc, charts["tam"], 5.5, "Exhibit 1: TAM / SAM / SOM Sizing (INR Crores)")

    section_head(doc, "4.  Competitive Landscape")
    comp = tables.get("competitive", {})
    rows_c = comp.get("rows", [])
    if rows_c:
        hdrs = ["Company","Stage","Funding","Core Product",
                "Rev Model","India","Moat","Weakness","Threat"]
        data = [[r.get("company",""), r.get("stage",""),
                 r.get("funding_raised",""), r.get("core_product",""),
                 r.get("revenue_model",""), r.get("india_focus",""),
                 r.get("moat",""), r.get("weakness",""),
                 r.get("threat_to_subject","")]
                for r in rows_c]
        styled_table(doc, hdrs, data)
    s = comp.get("competitive_summary","")
    if s: add_rich_text(doc, f"Analyst View: {s}")
    doc.add_paragraph()

    section_head(doc, "5.  Team Assessment")
    team = tables.get("team", {})
    members = team.get("members", [])
    if members:
        hdrs = ["Name","Title","Prior Employer","Prior Role",
                "Yrs Exp","Relevance","Education","Flag"]
        data = [[m.get("name",""), m.get("title",""),
                 m.get("previous_employer",""), m.get("previous_role",""),
                 m.get("years_experience",""), m.get("direct_relevance",""),
                 m.get("education",""), m.get("flag","")]
                for m in members]
        styled_table(doc, hdrs, data)
    for fld, lbl in [("team_verdict","Team Assessment"),
                     ("missing_hires","Missing Hires"),
                     ("founder_risk_reason","Founder Risk")]:
        v = team.get(fld,"")
        if v: add_rich_text(doc, f"{lbl}: {v}")
    doc.add_paragraph()

    section_head(doc, "6.  Unit Economics & Financial Analysis")
    fin = tables.get("financials", {})
    metrics = fin.get("metrics", [])
    if metrics:
        hdrs = ["Metric","Reported Value","Analyst Comment",
                "India Benchmark","Rating"]
        data = [[m.get("metric",""), m.get("reported_value",""),
                 m.get("analyst_comment",""), m.get("india_benchmark",""),
                 m.get("rating","")]
                for m in metrics]
        styled_table(doc, hdrs, data)
    if "fin" in charts:
        embed_img(doc, charts["fin"], 6.5, "Exhibit 2: Unit Economics Assessment Dashboard")
    for fld, lbl in [("headline_positive","Key Positive"),
                     ("headline_concern","Key Concern"),
                     ("burn_runway_comment","Burn and Runway")]:
        v = fin.get(fld,"")
        if v: add_rich_text(doc, f"{lbl}: {v}")
    doc.add_paragraph()

    section_head(doc, "7.  Key Risks & Mitigants")
    risk = tables.get("risks", {})
    risk_rows = risk.get("risks", [])
    if risk_rows:
        hdrs = ["ID","Category","Risk","Description","Prob","Impact",
                "Time","Mitigant","Residual","Deal Breaker"]
        data = [[r.get("id",""), r.get("category",""), r.get("title",""),
                 r.get("description",""), r.get("probability",""),
                 r.get("impact",""), r.get("time_horizon",""),
                 r.get("mitigant",""), r.get("residual_risk",""),
                 r.get("deal_breaker","")]
                for r in risk_rows]
        styled_table(doc, hdrs, data)
    if "risk" in charts:
        embed_img(doc, charts["risk"], 5.5, "Exhibit 3: Risk Matrix - Probability vs Impact")
    rs = risk.get("risk_summary","")
    if rs: add_rich_text(doc, f"Risk Summary: {rs}")
    doc.add_paragraph()

    doc.add_page_break()
    section_head(doc, "8.  Investment Recommendation")
    add_rich_text(doc, texts.get("recommendation",""))

    doc.add_paragraph()
    p = doc.add_paragraph("-" * 95)
    p.runs[0].font.size = Pt(8); p.runs[0].font.color.rgb = hex_rgb(GREY)
    disc = doc.add_paragraph(
        "DISCLAIMER: This memorandum was prepared using DealMemo AI as a "
        "first-draft analytical tool. All figures, assessments, and "
        "recommendations must be independently verified by a qualified "
        "investment professional before any capital commitment. This "
        "document does not constitute financial advice or a solicitation to invest."
    )
    disc.runs[0].font.size = Pt(8)
    disc.runs[0].font.color.rgb = hex_rgb(GREY)
    disc.runs[0].font.name = "Calibri"

    buf = BytesIO()
    doc.save(buf); buf.seek(0)
    return buf

# ════════════════════════════════════════════════════════════════════════════
#  STREAMLIT UI
# ════════════════════════════════════════════════════════════════════════════

st.set_page_config(
    page_title="DealMemo AI",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="collapsed"
)
st.markdown("""
<style>
.stButton button {
    background-color:#1F4E79;color:white;
    font-weight:600;border-radius:4px;border:none;
}
.stButton button:hover{background-color:#2E75B6;}
</style>
""", unsafe_allow_html=True)

st.title("📋 DealMemo AI")
st.subheader("Professional IC Memo Generator — Indian Startup Ecosystem")
st.markdown(
    "Upload a pitch deck and get a **full Investment Committee memo** "
    "— analyst-grade, industry-calibrated, IC-ready — in 4 minutes."
)
st.divider()

col1, col2 = st.columns(2)
with col1:
    company   = st.text_input("Company Name",
                  placeholder="e.g. Zomato, Groww, YourStartup")
    fund_type = st.selectbox("Memo Format", [
        "Angel Syndicate","Category 1 AIF","Category 2 AIF",
        "Family Office","Micro VC Fund","PE / Growth Equity"
    ])
with col2:
    pdf_file = st.file_uploader("Upload Pitch Deck (PDF)", type=["pdf"])
    if pdf_file:
        st.success(f"Uploaded: {pdf_file.name}")
    st.caption("Your deck is processed in memory only. Never stored.")

st.divider()

if st.button("Generate IC Memo", type="primary", use_container_width=True):
    if not company:
        st.error("Please enter the company name.")
    elif not pdf_file:
        st.error("Please upload a pitch deck PDF.")
    else:
        deck = extract_pdf(pdf_file)
        if not deck.strip():
            st.error("Could not extract text from PDF. Ensure it is not a scanned image-only file.")
            st.stop()

        bar    = st.progress(0)
        status = st.empty()
        STEPS  = 13
        step   = [0]

        def tick(msg):
            step[0] += 1
            bar.progress(min(step[0]/STEPS, 1.0))
            status.text(msg)

        texts  = {}
        tables = {}
        charts = {}

        tick("Analysing company profile...")
        profile = analyse_company_profile(deck, company)
        ctx     = profile_to_context(profile, company)
        industry = profile.get("primary_industry", "")
        stage    = profile.get("company_stage", "")
        listing  = profile.get("listing_status", "Private")
        st.info(f"Detected: {industry} | Stage: {stage} | Status: {listing} — calibrating all sections.")
        time.sleep(2)

        tick("Writing Executive Summary...")
        texts["exec"] = safe_ai_call([
            {"role":"system","content":SYSTEM},
            {"role":"user","content":prompt_exec(deck,company,fund_type,ctx)}
        ], max_tokens=1100)
        time.sleep(3)

        tick("Writing Company Overview...")
        texts["overview"] = safe_ai_call([
            {"role":"system","content":SYSTEM},
            {"role":"user","content":prompt_overview(deck,company,fund_type,ctx)}
        ], max_tokens=1000)
        time.sleep(3)

        tick("Writing Market Opportunity...")
        texts["market"] = safe_ai_call([
            {"role":"system","content":SYSTEM},
            {"role":"user","content":prompt_market(deck,company,fund_type,ctx)}
        ], max_tokens=1000)
        time.sleep(3)

        tick("Writing Investment Recommendation...")
        texts["recommendation"] = safe_ai_call([
            {"role":"system","content":SYSTEM},
            {"role":"user","content":prompt_recommendation(deck,company,fund_type,ctx)}
        ], max_tokens=1100)
        time.sleep(3)

        tick("Building Competitive Matrix...")
        raw = safe_ai_call([
            {"role":"system","content":"Return only valid JSON. No text."},
            {"role":"user","content":prompt_competitive(deck,company,profile)}
        ], expect_json=True, max_tokens=1400)
        try: tables["competitive"] = json.loads(raw)
        except Exception: tables["competitive"] = {"rows":[]}
        time.sleep(3)

        tick("Assessing Team...")
        raw = safe_ai_call([
            {"role":"system","content":"Return only valid JSON. No text."},
            {"role":"user","content":prompt_team(deck,company,profile)}
        ], expect_json=True, max_tokens=1400)
        try: tables["team"] = json.loads(raw)
        except Exception: tables["team"] = {"members":[]}
        time.sleep(3)

        tick("Analysing Unit Economics...")
        raw = safe_ai_call([
            {"role":"system","content":"Return only valid JSON. No text."},
            {"role":"user","content":prompt_financials(deck,company,fund_type,profile)}
        ], expect_json=True, max_tokens=1600)
        try: tables["financials"] = json.loads(raw)
        except Exception: tables["financials"] = {"metrics":[]}
        time.sleep(3)

        tick("Building Risk Matrix...")
        raw = safe_ai_call([
            {"role":"system","content":"Return only valid JSON. No text."},
            {"role":"user","content":prompt_risks(deck,company,fund_type,profile)}
        ], expect_json=True, max_tokens=1600)
        try: tables["risks"] = json.loads(raw)
        except Exception: tables["risks"] = {"risks":[]}
        time.sleep(2)

        tick("Generating TAM/SAM/SOM chart...")
        raw = safe_ai_call([
            {"role":"system","content":"Return only valid JSON. No text."},
            {"role":"user","content":prompt_tam(deck,company,profile)}
        ], expect_json=True, max_tokens=80)
        try:
            td = json.loads(raw)
            charts["tam"] = chart_tam(td if td.get("tam",0) > 0
                                      else {"tam":45000,"sam":10000,"som":2000}, company)
        except Exception:
            charts["tam"] = chart_tam({"tam":45000,"sam":10000,"som":2000}, company)

        tick("Generating Unit Economics chart...")
        if tables["financials"].get("metrics"):
            charts["fin"] = chart_financials(tables["financials"]["metrics"])

        tick("Generating Risk Heatmap...")
        if tables["risks"].get("risks"):
            charts["risk"] = chart_risk(tables["risks"]["risks"])

        bar.progress(1.0)
        status.text("IC Memo complete!")
        st.success(f"IC Memo for {company} is ready!")
        st.divider()

        t1,t2,t3,t4 = st.tabs(["Written Analysis","Tables","Charts","Risk Matrix"])
        with t1:
            for lbl, key in [
                ("Executive Summary","exec"),
                ("Company Overview","overview"),
                ("Market Opportunity","market"),
                ("Investment Recommendation","recommendation"),
            ]:
                with st.expander(lbl, expanded=False):
                    st.write(texts.get(key,""))
        with t2:
            st.markdown("**Competitive Landscape**")
            cr = tables["competitive"].get("rows",[])
            if cr: st.dataframe(cr, use_container_width=True)
            st.markdown("**Team Assessment**")
            tr = tables["team"].get("members",[])
            if tr: st.dataframe(tr, use_container_width=True)
            st.markdown("**Unit Economics**")
            mr = tables["financials"].get("metrics",[])
            if mr: st.dataframe(mr, use_container_width=True)
        with t3:
            if "tam" in charts:
                charts["tam"].seek(0)
                st.image(charts["tam"],caption="TAM/SAM/SOM",use_container_width=True)
            if "fin" in charts:
                charts["fin"].seek(0)
                st.image(charts["fin"],caption="Unit Economics",use_container_width=True)
        with t4:
            rr = tables["risks"].get("risks",[])
            if rr: st.dataframe(rr,use_container_width=True)
            if "risk" in charts:
                charts["risk"].seek(0)
                st.image(charts["risk"],caption="Risk Heatmap",use_container_width=True)

        st.divider()

        for k in ["tam","fin","risk"]:
            if k in charts: charts[k].seek(0)

        doc_buf = build_memo(texts, tables, charts, company, fund_type, profile)

        st.download_button(
            label="Download Full IC Memo - Word Document (.docx)",
            data=doc_buf,
            file_name=f"IC_Memo_{company.replace(' ','_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        st.caption("Review all sections before presenting to IC. Verify all figures independently.")
